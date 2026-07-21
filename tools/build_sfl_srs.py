from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("tmp/srs_build/SFL_SRS.docx")


PHASE1_SYSTEMS = [
    {
        "id": "S152",
        "name": "Computer-Aided Facility Management (CAFM) / IWMS",
        "short": "CAFM/IWMS",
        "platform": "SFL.IFIMP",
        "unit": "Building & Infrastructure Unit",
        "purpose": "Facilities master data, buildings, rooms, zones, service control and examination readiness.",
        "objects": "Site, building, floor, room, zone, facility asset, readiness profile, service request.",
        "integrations": "IAM/OIDC, HRMS staff directory, NBES examination centre schedules, CMMS, room booking, asset reference, audit/evidence, reporting.",
        "dashboard": "Facility readiness, open service requests, unavailable rooms, site compliance exceptions and examination readiness risk.",
    },
    {
        "id": "S153",
        "name": "Computerized Maintenance Management System (CMMS)",
        "short": "CMMS",
        "platform": "SFL.IFIMP",
        "unit": "Building & Infrastructure Unit",
        "purpose": "Work orders, preventive maintenance, faults, vendor SLA tracking, parts and evidence closure.",
        "objects": "Work order, fault report, preventive schedule, vendor assignment, SLA timer, closure evidence.",
        "integrations": "CAFM/IWMS, BMS/fire/life-safety events where available, procurement/vendor master, notifications, audit/evidence, reporting.",
        "dashboard": "Open work orders, overdue SLAs, blocked readiness items, repeated faults, vendor performance and closure evidence gaps.",
    },
    {
        "id": "S159",
        "name": "Room and Resource Booking System",
        "short": "Room Booking",
        "platform": "SFL.IFIMP",
        "unit": "Building & Infrastructure Unit",
        "purpose": "Room bookings, lecture halls, moot courtrooms, resources, conflict checks, setup tasks and no-show enforcement.",
        "objects": "Booking, room, resource, approval, setup task, readiness hold, no-show record.",
        "integrations": "CAFM/IWMS, calendar provider, LMS/NBES schedules, security, logistics, notifications, audit/evidence.",
        "dashboard": "Utilization, conflicts, pending approvals, setup readiness, no-shows and examination-mode locks.",
    },
    {
        "id": "S160",
        "name": "Visitor Management System (VMS)",
        "short": "Visitor Management",
        "platform": "SFL.SSEMP",
        "unit": "Health, Safety & Security Unit",
        "purpose": "Pre-registration, host approval, badge printing, check-in/out, NDA capture, watchlist screening and evacuation roll-call.",
        "objects": "Visitor, visit, host, badge, access zone, watchlist result, roll-call status.",
        "integrations": "IAM/OIDC, HRMS hosts, access control, emergency notification, watchlist source where approved, audit/evidence, reporting.",
        "dashboard": "On-site visitors, overdue check-outs, denied entries, roll-call completion and host approval backlog.",
    },
    {
        "id": "S160a",
        "name": "Physical Access Control Integration",
        "short": "Access Control",
        "platform": "SFL.SSEMP",
        "unit": "Health, Safety & Security Unit",
        "purpose": "Door access, card/biometric events, access rules, overrides and joiner-mover-leaver support.",
        "objects": "Access identity, credential, reader, door, access zone, access event, override.",
        "integrations": "IAM/OIDC, HRMS joiner-mover-leaver events, visitor management, device controllers, audit/evidence, SIEM/SOC reporting.",
        "dashboard": "Denied access, forced door events, reader health, active overrides, credential exceptions and restricted-zone exposure.",
    },
    {
        "id": "S161",
        "name": "CCTV / Video Management System Integration",
        "short": "CCTV/VMS",
        "platform": "SFL.SSEMP",
        "unit": "Health, Safety & Security Unit",
        "purpose": "Camera health, incident linkage, evidence request, controlled footage export and retention governance.",
        "objects": "Camera, video system, camera health, evidence request, export approval, retention hold.",
        "integrations": "VMS/NVR APIs, incident management, access/alarm events, data protection/DPIA, audit/evidence, investigations.",
        "dashboard": "Camera offline events, recording gaps, pending evidence requests, export approvals and retention holds.",
    },
    {
        "id": "S162",
        "name": "Intrusion Detection and Alarm Monitoring",
        "short": "Intrusion Alarm",
        "platform": "SFL.SSEMP",
        "unit": "Health, Safety & Security Unit",
        "purpose": "Intrusion alerts, restricted-zone alarms, SOC queue, acknowledgement, escalation and restoration.",
        "objects": "Alarm panel, zone, alarm event, acknowledgement, escalation, restoration.",
        "integrations": "Alarm panels, CCTV/VMS, access control, emergency notification, incident management, audit/evidence.",
        "dashboard": "Open alarms, acknowledgement time, repeated zones, tamper events, restoration lag and escalation status.",
    },
    {
        "id": "S162a",
        "name": "Fire-Safety and Life-Safety Monitoring",
        "short": "Life Safety",
        "platform": "SFL.SSEMP",
        "unit": "Health, Safety & Security Unit",
        "purpose": "Fire alarms, smoke/panic events, inspection tracking and emergency response triggers. Vendor life-safety systems remain authoritative for actuation.",
        "objects": "Life-safety panel, detector, panic point, inspection, event, response trigger.",
        "integrations": "Fire panels/BMS bridge, CMMS, emergency notification, evacuation drill records, audit/evidence, reporting.",
        "dashboard": "Fire/panic events, panel faults, overdue inspections, emergency fast-lane events and certification exceptions.",
    },
    {
        "id": "S163",
        "name": "Health and Safety Incident / Near-Miss Reporting",
        "short": "HSE Incidents",
        "platform": "SFL.SSEMP",
        "unit": "Health, Safety & Security Unit",
        "purpose": "HSE incident reporting, near-miss capture, root cause analysis, corrective actions and statutory reporting support.",
        "objects": "Incident, near miss, injury, RCA, corrective action, witness statement, evidence attachment.",
        "integrations": "Unified portal/intranet, HRMS/P&C where staff are involved, insurance/claims, compliance/audit, notifications, reporting.",
        "dashboard": "Open incidents, CAPA aging, severity trends, repeat locations, reportable cases and overdue reviews.",
    },
    {
        "id": "S166",
        "name": "Fleet and Vehicle Management",
        "short": "Fleet",
        "platform": "SFL.FTLMP",
        "unit": "Transportation & Logistics Unit",
        "purpose": "Fleet register, vehicle compliance, service status, readiness, assignment and driver eligibility.",
        "objects": "Vehicle, driver, assignment, compliance document, service status, availability.",
        "integrations": "HRMS driver/staff records, CMMS/service vendors, fuel/logbooks, dispatch, GPS/telematics-ready adapter, audit/evidence.",
        "dashboard": "Vehicle availability, expired compliance, service due, assignment conflicts and readiness blockers.",
    },
    {
        "id": "S168_fuel",
        "name": "Fuel Management and Driver Logbooks",
        "short": "Fuel & Logbooks",
        "platform": "SFL.FTLMP",
        "unit": "Transportation & Logistics Unit",
        "purpose": "Fuel transactions, driver logbooks, odometer capture, reconciliation, exception review and anomaly detection.",
        "objects": "Fuel transaction, fuel limit, receipt, odometer reading, driver logbook, anomaly case.",
        "integrations": "Fleet, fuel card/provider data, finance, audit/evidence, notifications, reporting.",
        "dashboard": "Fuel spend, variances, missing receipts, abnormal consumption, unresolved anomalies and driver logbook compliance.",
    },
    {
        "id": "S171",
        "name": "Mailroom / Courier and Dispatch Tracking",
        "short": "Courier & Dispatch",
        "platform": "SFL.FTLMP",
        "unit": "Transportation & Logistics Unit",
        "purpose": "Courier item register, sealed dispatches, manifests, receipt confirmation, exceptions and chain-of-custody.",
        "objects": "Courier item, dispatch manifest, custody handoff, delivery receipt, exception, seal reference.",
        "integrations": "Examination operations/NBES, document management, fleet, notifications, audit/evidence, reporting.",
        "dashboard": "Items in transit, overdue receipts, broken seals, custody gaps, dispatch exceptions and examination material movements.",
    },
    {
        "id": "S174",
        "name": "Emergency Mass Notification System",
        "short": "Emergency Notification",
        "platform": "SFL.SSEMP",
        "unit": "Health, Safety & Security Unit",
        "purpose": "Emergency templates, zone selection, channel activation, break-glass send, delivery receipts and acknowledgement tracking.",
        "objects": "Notification template, emergency scenario, recipient zone, channel, send event, acknowledgement.",
        "integrations": "SMS/email/push/voice/siren/signage providers, visitor management, access control, alarm/life-safety events, audit/evidence, reporting.",
        "dashboard": "Activation time, channel delivery, acknowledgements, failed recipients, break-glass usage and drill performance.",
    },
]


FULL_CLUSTER = [
    ("S152", "Computer-Aided Facility Management (CAFM) / IWMS", "Fast-Track"),
    ("S153", "Computerized Maintenance Management System (CMMS)", "Fast-Track"),
    ("S156", "Building Management System (BMS) / IoT", "Phase 2"),
    ("S157", "Energy & Sustainability Monitoring", "Phase 2"),
    ("S158", "Space Planning & Move Management", "Phase 2"),
    ("S159", "Room and Resource Booking System", "Fast-Track"),
    ("S160", "Visitor Management System (VMS)", "Fast-Track"),
    ("S160a", "Physical Access Control Integration", "Fast-Track"),
    ("S161", "CCTV / Video Management System (VMS)", "Fast-Track"),
    ("S162", "Intrusion Detection and Alarm Monitoring", "Fast-Track"),
    ("S162a", "Fire-Safety and Life-Safety Monitoring", "Fast-Track"),
    ("S163", "Health and Safety Incident and Near-Miss Reporting", "Fast-Track"),
    ("S164", "Permit-to-Work / Hot-Work Authorization", "Phase 2"),
    ("S165", "Risk Assessment Library (HSE)", "Phase 2"),
    ("S166", "Fleet and Vehicle Management", "Fast-Track"),
    ("S167", "GPS / Telematics Tracking", "Phase 2"),
    ("S168", "Asset Tagging / RFID / Barcode Inventory", "Phase 2"),
    ("S168_fuel", "Fuel Management and Driver Logbooks", "Fast-Track"),
    ("S168a", "Trip / Driver-Booking Portal", "Phase 2"),
    ("S169", "Cleaning and Janitorial Schedule Management", "Phase 2"),
    ("S170", "Pest-Control and Hygiene Audit Tracker", "Phase 3"),
    ("S171", "Mailroom / Courier and Dispatch Tracking", "Fast-Track"),
    ("S172", "Catering and Cafeteria Management", "Phase 3"),
    ("S173", "Event Logistics and Set-Up Workflow", "Phase 2"),
    ("S174", "Emergency Mass-Notification System", "Fast-Track"),
    ("S175", "Crisis and Evacuation Drill Management", "Phase 2"),
    ("S176", "Construction Project Management", "Phase 2"),
    ("S177", "Lease and Tenancy Management", "Phase 3"),
    ("S178", "Waste Management and Recycling Tracking", "Phase 3"),
    ("S179", "Lost-and-Found Register", "Phase 3"),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Calibri"
    r.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, True)
        set_cell_shading(hdr[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    if widths:
        set_table_widths(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_requirement(
    doc: Document,
    ref: str,
    title: str,
    actor: str,
    goal: str,
    benefit: str,
    requirements: list[str],
    fields: list[str],
    validation: list[str],
    workflow: list[str],
    errors: list[str],
    acceptance: list[str],
) -> None:
    doc.add_heading(f"{ref}: {title}", level=3)
    doc.add_paragraph("User Story", style="SFL Label")
    doc.add_paragraph(f"As a {actor}, I want to {goal}, so that {benefit}.")
    doc.add_paragraph("Requirements", style="SFL Label")
    add_bullets(doc, requirements)
    if fields:
        doc.add_paragraph("System-managed fields", style="SFL Label")
        add_bullets(doc, fields)
    doc.add_paragraph("Validation", style="SFL Label")
    add_bullets(doc, validation)
    doc.add_paragraph("Workflow", style="SFL Label")
    add_numbered(doc, workflow)
    doc.add_paragraph("Error States", style="SFL Label")
    add_bullets(doc, errors)
    doc.add_paragraph("Acceptance Criteria", style="SFL Label")
    add_bullets(doc, acceptance)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    if "SFL Label" not in styles:
        label = styles.add_style("SFL Label", 1)
        label.font.name = "Calibri"
        label.font.size = Pt(10)
        label.font.bold = True
        label.font.color.rgb = RGBColor.from_string("1F4D78")
        label.paragraph_format.space_before = Pt(4)
        label.paragraph_format.space_after = Pt(2)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("CLET SFL SRS v1.0")
    return doc


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("COUNCIL FOR LEGAL EDUCATION AND TRAINING (CLET)")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor.from_string("0B2545")

    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Cluster 9: Facilities, Safety & Logistics")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor.from_string("0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Software Requirements Specification (SRS)")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string("2E74B5")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Safety, Facilities and Logistics Phase 1 Platform - 13 Fast-Track Systems").bold = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Implementation target: Java Spring Boot modular platform")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Version 1.0 - July 2026")
    doc.add_paragraph()
    doc.add_paragraph("Legislative and governance basis: CLET Comprehensive Digital System Mapping v2; Cluster 9 FSL System Architecture Document; Ghana Data Protection Act, 2012 (Act 843); applicable health, safety, procurement, evidence, audit and institutional governance obligations.")
    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    doc.add_heading("Document Control", level=1)
    add_table(doc, ["Item", "Value"], [
        ["Document Title", "CLET Cluster 9 Safety, Facilities and Logistics Software Requirements Specification"],
        ["Short Name", "SFL_SRS"],
        ["Version", "1.0"],
        ["Date", "July 2026"],
        ["Prepared For", "Council for Legal Education and Training (CLET)"],
        ["Primary Source of Truth", "CLET_Comprehensive_Digital_System_Mapping_v2"],
        ["Derived From", "CLET_Cluster9_FSL_System_Architecture_Document_FULLY_INTEGRATED and SFL_Phase1_System_Architecture_Implementation_Guide_v2"],
        ["Implementation Alignment", "Java Spring Boot modular monolith with bounded contexts and extraction-ready contracts"],
    ], [1.8, 4.7])

    doc.add_heading("Distribution Control", level=1)
    add_table(doc, ["Audience", "Purpose"], [
        ["CLET Executive / Management", "Approve scope, priorities, operational governance and readiness gates."],
        ["Facilities and Logistics Directorate", "Validate operational workflows, roles, KPIs and service ownership."],
        ["DTI / Architecture / Development Team", "Scaffold and implement the Spring Boot platform against approved requirements."],
        ["Security, HSE and Emergency Operations", "Validate safety, security, evidence and emergency response controls."],
        ["Audit, Compliance and Data Protection", "Confirm auditability, privacy, retention and evidentiary controls."],
    ], [2.0, 4.5])

    doc.add_heading("Systems", level=1)
    rows = [[str(i + 1), s["id"], s["name"], "Phase 1 / Fast-Track", s["platform"]] for i, s in enumerate(PHASE1_SYSTEMS)]
    add_table(doc, ["No.", "System ID", "System Name", "Phase", "Platform"], rows, [0.45, 0.8, 3.25, 1.05, 0.95])

    doc.add_heading("Revision History", level=1)
    add_table(doc, ["Version", "Date", "Description", "Author"], [
        ["1.0", "July 2026", "Initial SFL SRS generated from Cluster 9 source mapping, integrated architecture document and phase 1 implementation guide.", "Codex / CLET project workspace"],
    ], [0.7, 0.9, 4.1, 0.8])

    doc.add_heading("Approvals", level=1)
    add_table(doc, ["Role", "Name", "Signature", "Date"], [
        ["Director, Facilities and Logistics", "", "", ""],
        ["Director, DTI", "", "", ""],
        ["Director-General / Delegate", "", "", ""],
        ["Audit / Compliance Representative", "", "", ""],
    ], [2.4, 1.4, 1.4, 1.3])

    doc.add_heading("References", level=1)
    add_bullets(doc, [
        "CLET_Comprehensive_Digital_System_Mapping_v2 - authoritative enterprise system mapping and Cluster 9 phase classification.",
        "CLET_Cluster9_FSL_System_Architecture_Document_FULLY_INTEGRATED - integrated functional, non-functional, architecture, workflow, role and acceptance baseline.",
        "SFL_Phase1_System_Architecture_Implementation_Guide_v2 - implementation architecture, release sequence, edge/core model, audit, integration security and Spring Boot realignment basis.",
        "C7_HRM_SRS - formatting and SRS structure reference.",
        "Ghana Data Protection Act, 2012 (Act 843) - personal-data, biometric, visitor, CCTV and evidence governance baseline.",
    ])

    doc.add_heading("Table of Contents", level=1)
    toc = [
        "1. Introduction",
        "2. Overall Description",
        "3. Software Requirement Specifications",
        "21. Data Requirements",
        "22. External Interface Requirements",
        "23. Non-Functional Requirements",
        "Appendix A. Phase and Source Traceability",
    ]
    add_bullets(doc, toc)
    doc.add_page_break()


def add_intro(doc: Document) -> None:
    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Purpose", level=2)
    doc.add_paragraph(
        "This document defines the functional and non-functional software requirements for the Safety, Facilities and Logistics (SFL) Phase 1 platform of CLET. It translates the authoritative Cluster 9 system mapping into an implementation-ready SRS that can guide Spring Boot scaffolding, domain modelling, API design, testing, commissioning and go-live readiness."
    )
    doc.add_heading("1.2 Scope", level=2)
    doc.add_paragraph(
        "The Phase 1 scope contains 13 fast-track systems from Cluster 9. These systems are implemented as a coordinated SFL platform rather than 13 disconnected applications. The platform includes facilities operations, maintenance, room booking, visitor management, physical access control, CCTV integration, intrusion alarms, fire/life-safety monitoring, HSE incident reporting, fleet management, fuel and driver logbooks, courier/dispatch tracking and emergency mass notification."
    )
    doc.add_paragraph(
        "SFL.AVAMP-Lite is included as a supporting reference capability for asset/device identifiers, custody/location references and device-to-location mapping. Full Asset Tagging / RFID / Barcode Inventory (S168) remains Phase 2 in the parent mapping and is not treated as a full Phase 1 business system."
    )
    doc.add_heading("1.3 Intended Audience", level=2)
    add_bullets(doc, [
        "CLET management and operational owners who approve scope and operational readiness.",
        "Facilities, Security, HSE, Fleet, Logistics and Emergency Operations teams who validate workflows.",
        "Java Spring Boot developers, architects and testers who implement the platform.",
        "Audit, compliance and data protection teams who validate logging, evidence, retention and privacy controls.",
        "Procurement and vendor-integration stakeholders who provide device, CCTV, access-control, alarm, notification and fleet system inputs.",
    ])
    doc.add_heading("1.4 Definitions, Acronyms and Abbreviations", level=2)
    add_table(doc, ["Term", "Definition"], [
        ["SFL", "Safety, Facilities and Logistics platform for CLET Cluster 9."],
        ["IFIMP", "Integrated Facilities and Infrastructure Management Platform."],
        ["SSEMP", "Safety, Security and Emergency Management Platform."],
        ["FTLMP", "Fleet, Transport and Logistics Management Platform."],
        ["AVAMP-Lite", "Phase 1 asset/device reference layer supporting location, custody and device identifiers."],
        ["CAFM/IWMS", "Computer-Aided Facilities Management / Integrated Workplace Management System."],
        ["CMMS", "Computerized Maintenance Management System."],
        ["SOC", "Security Operations Centre or security monitoring function."],
        ["HSE", "Health, Safety and Environment."],
        ["CAPA", "Corrective and Preventive Action."],
        ["RTO/RPO", "Recovery time objective / recovery point objective."],
    ], [1.35, 5.15])
    doc.add_heading("1.5 References", level=2)
    doc.add_paragraph("The references listed in the front matter are normative for this SRS. Where source documents conflict, the parent CLET Comprehensive Digital System Mapping v2 is treated as the source of truth for system identity, cluster ownership and phase classification.")
    doc.add_heading("1.6 Revision History", level=2)
    doc.add_paragraph("Version 1.0 establishes the SFL Phase 1 SRS baseline for implementation realignment from the earlier .NET architecture notes to a Java Spring Boot delivery stack.")


def add_overall(doc: Document) -> None:
    doc.add_heading("2. Overall Description", level=1)
    doc.add_heading("2.1 Product Perspective", level=2)
    doc.add_paragraph(
        "The SFL platform is the operational control layer for CLET facilities, safety, security, fleet and logistics operations. It integrates with identity, HRMS, examination operations, finance/procurement, document management, analytics, SIEM/SOC/NOC and communication systems through controlled APIs, signed webhooks, event contracts and vendor adapters. SFL owns operational records and workflows for Cluster 9 but does not replace enterprise identity, payment, procurement, HR or examination platforms."
    )
    doc.add_paragraph(
        "The implementation baseline is a Spring Boot modular monolith with bounded contexts aligned to SFL.IFIMP, SFL.SSEMP, SFL.FTLMP and SFL.AVAMP-Lite. Modules communicate through contracts, events and application ports. This preserves the implementation guide's extraction-ready architecture while aligning the current build to Java Spring Boot."
    )
    doc.add_heading("2.2 Product Functions - Overview", level=2)
    add_bullets(doc, [
        "Maintain authoritative site, building, room, zone, vehicle, device and logistics reference records.",
        "Manage facility faults, work orders, preventive maintenance and readiness evidence.",
        "Manage room and resource booking with conflict prevention, approval, setup tasks and examination-mode locks.",
        "Manage visitors, access-control events, CCTV evidence requests, alarms, fire/life-safety events and HSE incidents.",
        "Manage fleet availability, driver assignments, fuel transactions, driver logbooks, courier items and sealed dispatches.",
        "Send and audit emergency notifications across approved channels, with break-glass behaviour for declared emergencies.",
        "Provide dashboards, operational KPIs, evidence trails, audit records, integration health and readiness reporting.",
    ])
    doc.add_heading("2.3 User Classes and Characteristics", level=2)
    add_table(doc, ["Role", "Description"], [
        ["Facilities Director / Manager", "Owns facilities policy, readiness, escalations, dashboards and approval of critical operational exceptions."],
        ["Facilities Officer", "Maintains sites, rooms, resources, work orders, preventive maintenance and readiness checks."],
        ["Maintenance Technician / Vendor", "Receives assigned work, records progress, uploads evidence and closes assigned tasks subject to verification."],
        ["Room Requester / Host", "Requests rooms/resources, receives booking decisions and manages attendance/no-show obligations."],
        ["Security Director / SOC Operator", "Monitors visitors, access events, CCTV, intrusion alarms, fire/life-safety alerts and emergency queues."],
        ["Visitor Desk Officer", "Pre-registers, checks in/out, badges and supports visitor roll-call."],
        ["HSE Officer", "Reviews incidents, runs RCA/CAPA workflows, tracks corrective actions and supports statutory reporting."],
        ["Fleet / Logistics Officer", "Manages vehicles, drivers, fuel, logbooks, courier movements, dispatch manifests and custody events."],
        ["Emergency Coordinator", "Activates approved emergency notifications, manages acknowledgement and monitors drills/incidents."],
        ["Auditor / Compliance / Data Protection Officer", "Reviews audit trails, evidence exports, retention, access logs and data protection controls."],
        ["System Administrator", "Configures users, roles, integration adapters, sites, zones, thresholds, SLAs, retention and system parameters."],
    ], [2.05, 4.45])
    doc.add_heading("2.4 Status Matrix", level=2)
    add_table(doc, ["No.", "Action", "Internal Status", "User-facing Status", "Manual Status Change"], [
        ["1", "Facility fault reported", "New", "Submitted", "No"],
        ["2", "Work order assigned", "Assigned", "Assigned to technician/vendor", "No"],
        ["3", "Work completed", "Pending Verification", "Awaiting verification", "No"],
        ["4", "Work order verified", "Closed", "Completed", "No"],
        ["5", "Room booking requested", "Pending Approval", "Pending approval", "No"],
        ["6", "Booking approved", "Approved / Setup Pending", "Approved", "No"],
        ["7", "Visitor pre-registered", "Pre-Registered", "Invitation sent", "No"],
        ["8", "Visitor checked in", "On Site", "Checked in", "No"],
        ["9", "Access event denied", "Access Denied", "Denied / under review", "No"],
        ["10", "Security/life-safety alert raised", "Open Alert", "Response in progress", "No"],
        ["11", "HSE incident reported", "Triage", "Submitted for review", "No"],
        ["12", "Emergency notification sent", "Sent / Awaiting Ack", "Alert sent", "No"],
        ["13", "Vehicle assigned", "Assigned", "Vehicle assigned", "No"],
        ["14", "Fuel anomaly detected", "Under Review", "Under review", "No"],
        ["15", "Dispatch receipt confirmed", "Delivered", "Delivered", "No"],
    ], [0.45, 1.8, 1.35, 1.9, 1.0])
    doc.add_heading("2.5 Operating Environment", level=2)
    add_bullets(doc, [
        "Web-based Spring Boot platform with REST APIs and browser/mobile-responsive portal screens.",
        "PostgreSQL operational store with schema per module and row-level security for site-scoped and sensitive records.",
        "Kafka-compatible event bus or pluggable event adapter for integration events and outbox delivery.",
        "Redis-compatible cache or pluggable cache adapter for device state, permission snapshots and dashboard acceleration.",
        "OIDC/OAuth2 identity provider integration through an adapter; Keycloak may be used but must not leak into domain code.",
        "Dockerized development, staging and production deployment with OpenTelemetry logs, metrics and traces.",
        "Core tier at HQ/DR and edge tier per examination centre where local operation is required during WAN loss.",
    ])
    doc.add_heading("2.6 Design and Implementation Constraints", level=2)
    add_bullets(doc, [
        "SFL must not implement authentication, MFA, login, logout or password reset. These are delegated to the enterprise identity provider.",
        "Spring Boot modules must communicate through public contracts, application ports and integration events, not cross-module database joins.",
        "No cross-schema foreign keys are permitted between SFL modules. Cross-module consistency is eventual and event-driven.",
        "Inbound device/vendor events must be authenticated, schema-validated, idempotent and stored before domain processing.",
        "Certified life-safety systems remain authoritative for alarm actuation. SFL observes, records, notifies and governs; it must not sit in the certified actuation path.",
        "SLA thresholds, zones, severities, fuel limits, readiness checklists, escalation rules and retention policies must be configurable without code redeployment.",
        "Every operational record must be site-scoped and auditable."
    ])
    doc.add_heading("2.7 Assumptions and Dependencies", level=2)
    doc.add_paragraph("Assumptions:")
    add_bullets(doc, [
        "CLET confirms site, building, room, vehicle, user-role and operational-owner master data before production load.",
        "Identity, HRMS, examination operations and vendor systems expose agreed APIs, files, webhooks or event feeds before integration testing.",
        "Device suppliers provide stable identifiers for cameras, access readers, alarm panels, fire/life-safety panels, vehicles and dispatch devices.",
        "Data retention and evidence-release rules are approved by management, audit and data-protection stakeholders before go-live.",
    ])
    doc.add_paragraph("Dependencies:")
    add_bullets(doc, [
        "Enterprise identity provider and user/role directory.",
        "HRMS for staff, host, driver and joiner-mover-leaver data.",
        "NBES/examination operations for centre schedules, examination mode and examination dispatch context.",
        "Vendor systems for CCTV, access control, alarms, fire/life-safety, SMS/email/voice/push/signage and fleet/fuel data.",
        "Audit/evidence store, SIEM/SOC/NOC monitoring, document management and analytics/reporting services.",
    ])
    doc.add_heading("2.8 System Use Case Overview", level=2)
    add_bullets(doc, [
        "Facilities officer reports or receives a fault, assigns a work order, tracks evidence and closes after verification.",
        "Room requester books a room/resource; the platform checks conflicts, approvals, setup tasks and examination-mode restrictions.",
        "Visitor desk officer registers and checks in a visitor; emergency coordinator uses the visitor population for roll-call.",
        "SOC operator receives access, CCTV, intrusion, fire/life-safety and panic events; triages, links evidence and escalates.",
        "HSE officer records incident details, performs RCA, assigns CAPA and monitors overdue corrective actions.",
        "Fleet/logistics officer assigns vehicles, reviews driver logbooks, reconciles fuel and tracks courier/dispatch chain-of-custody.",
        "Emergency coordinator sends approved or break-glass notifications and monitors delivery/acknowledgement by site and channel.",
    ])


def add_system_requirements(doc: Document) -> None:
    doc.add_heading("3. Software Requirement Specifications", level=1)
    doc.add_paragraph(
        "The requirements below use the C7 HRM SRS pattern and convert the Cluster 9 Phase 1 mapping into implementation-ready requirements. Requirement IDs use SRS-SFL-[SystemID]-[Number]."
    )

    actor_map = {
        "SFL.IFIMP": "Facilities Officer",
        "SFL.SSEMP": "Security or HSE Officer",
        "SFL.FTLMP": "Fleet or Logistics Officer",
    }

    for system in PHASE1_SYSTEMS:
        sid = system["id"].replace("_", "")
        actor = actor_map[system["platform"]]
        doc.add_heading(f"MODULE: {system['name']} ({system['id']})", level=2)
        doc.add_paragraph(f"Source mapping phase: Fast-Track. Platform: {system['platform']}. Operational owner: F&L | {system['unit']}. Purpose: {system['purpose']}")

        add_requirement(
            doc,
            f"SRS-SFL-{sid}-01",
            f"Maintain {system['short']} Operational Records",
            actor,
            f"create and maintain the operational records for {system['short']}",
            f"CLET has an accurate, site-scoped source of truth for {system['short'].lower()} operations",
            [
                f"The system shall maintain records for {system['objects']}.",
                "Each record shall be linked to an authorised site, building, room, zone, vehicle or device reference as applicable.",
                "Users shall only see and update records inside their assigned site scopes and roles.",
                "Records shall support active, inactive, suspended and archived lifecycle states where applicable.",
            ],
            [
                "Record UUID/ULID, site scope, created by/date, last modified by/date, version, source channel and audit correlation ID.",
            ],
            [
                "A record cannot be saved without the required site scope and operational owner.",
                "Duplicate active identifiers are blocked within the same site and object type.",
                "Sensitive fields are masked from roles without explicit permission.",
            ],
            [
                "Authorised user opens the relevant module.",
                "User creates or updates the operational record.",
                "System validates site scope, duplicates and required fields.",
                "System saves the record, writes audit evidence and publishes any required change event.",
            ],
            [
                'Duplicate Identifier - "An active record with this identifier already exists for this site."',
                'Missing Site Scope - "Select a valid CLET site before saving this record."',
                'Unauthorized Scope - "You are not authorised to access this site or record."',
            ],
            [
                "Given an authorised user creates a valid record, when the record is saved, then the system persists it with a unique identifier and audit trail.",
                "Given a user attempts to create a duplicate active identifier, when they save, then the system blocks the action and shows the duplicate error.",
                "Given a user lacks the required site scope, when they attempt to view the record, then the system denies access.",
            ],
        )

        add_requirement(
            doc,
            f"SRS-SFL-{sid}-02",
            f"Execute {system['short']} Workflow",
            actor,
            f"initiate, assign, escalate and close {system['short']} workflow items",
            "operational work is controlled through transparent queues, SLAs and accountable ownership",
            [
                "The system shall support workflow creation, assignment, reassignment, escalation, hold, cancellation and closure.",
                "The system shall calculate SLA timers from configurable priority, severity, site, operating mode and workflow type rules.",
                "The system shall notify responsible users when work is assigned, overdue, escalated or blocked.",
                "The system shall retain all workflow transitions and comments in the audit trail.",
            ],
            [
                "Workflow ID, status, priority, severity, assignee, SLA due date, escalation level, related record, closure reason and closure timestamp.",
            ],
            [
                "A workflow cannot be closed without required evidence or closure reason.",
                "Only authorised roles may approve, override, cancel or reopen workflow items.",
                "Escalation rules must be evaluated using the runtime configuration active at the time of evaluation.",
            ],
            [
                "User raises or receives a workflow item.",
                "System assigns the item based on configured role, site and queue rules.",
                "Assigned user records progress, evidence and comments.",
                "System escalates if SLA or risk thresholds are breached.",
                "Authorised user closes or verifies closure; system publishes completion events.",
            ],
            [
                'Closure Evidence Missing - "Required evidence must be attached before closure."',
                'SLA Breach - "This item has breached its configured SLA and has been escalated."',
                'Unauthorized Approval - "You do not have permission to approve this workflow transition."',
            ],
            [
                "Given a workflow item is assigned, when the assignee updates progress, then the system records the transition with actor and timestamp.",
                "Given required closure evidence is missing, when closure is attempted, then the system blocks closure.",
                "Given an SLA threshold is breached, when the scheduled evaluation runs, then the system escalates the item and notifies the configured role.",
            ],
        )

        add_requirement(
            doc,
            f"SRS-SFL-{sid}-03",
            f"Capture Evidence and Audit Trail for {system['short']}",
            "Auditor or Compliance Officer",
            f"review evidence, exports and audit events for {system['short']}",
            "CLET can prove who did what, when, why and with which evidence",
            [
                "The system shall capture actor, timestamp, before/after values, source channel and correlation ID for all state-changing actions.",
                "The system shall store evidence metadata including file reference, hash, uploader, related workflow, retention class and access history.",
                "Evidence access and export shall require role permission, justification and audit logging.",
                "Audit records shall be append-only and tamper-evident using a hash-chain or equivalent control.",
            ],
            [
                "Audit ID, previous hash, record hash, evidence ID, retention class, export approval, export recipient and legal hold flag.",
            ],
            [
                "Evidence cannot be exported without an approved reason and authorised role.",
                "Retention class is mandatory for CCTV, visitor, biometric, incident and dispatch evidence.",
                "Audit records cannot be modified or deleted by normal application roles.",
            ],
            [
                "User performs an operational action or uploads evidence.",
                "System writes the operational record and audit entry in the same unit of work where possible.",
                "System computes evidence hash and links metadata to the workflow.",
                "Authorised auditor searches, reviews and exports evidence where approved.",
            ],
            [
                'Export Not Approved - "Evidence export requires approval and a recorded reason."',
                'Retention Class Missing - "Select a retention class before saving this evidence."',
                'Audit Chain Failure - "Audit integrity check failed. Escalate to compliance and security."',
            ],
            [
                "Given evidence is uploaded, when the upload completes, then the system stores metadata, hash and audit reference.",
                "Given a user without export permission attempts export, when they submit the request, then the system blocks it.",
                "Given an audit chain replay detects tampering, when the integrity check runs, then the system raises a critical compliance alert.",
            ],
        )

        add_requirement(
            doc,
            f"SRS-SFL-{sid}-04",
            f"Integrate {system['short']} with Related Systems",
            "System Administrator",
            f"configure and monitor integrations for {system['short']}",
            "SFL exchanges data with CLET enterprise and vendor systems without direct database coupling",
            [
                f"The system shall integrate with: {system['integrations']}.",
                "All inbound webhooks, callbacks and device events shall verify HMAC signature or mTLS certificate, source allowlist and schema validity before processing.",
                "Integration messages shall carry idempotency keys and correlation IDs.",
                "Failed integration deliveries shall be retried and surfaced on an integration-health dashboard.",
            ],
            [
                "Integration message ID, idempotency key, payload schema version, source system, processing status, retry count, error reason and dead-letter reference.",
            ],
            [
                "Unsigned, untrusted or schema-invalid payloads must be rejected before reaching domain commands.",
                "No vendor adapter may write directly into operational module tables.",
                "Consumers must be idempotent because delivery is at-least-once.",
            ],
            [
                "External system sends or receives an integration message.",
                "Integration adapter authenticates, validates and stores the message envelope.",
                "Application command or event handler processes the payload idempotently.",
                "System records success, retry or dead-letter outcome and exposes integration health.",
            ],
            [
                'Invalid Signature - "Integration message rejected: signature verification failed."',
                'Schema Validation Failed - "Integration message rejected: payload does not match registered schema."',
                'Duplicate Message - "Duplicate integration message received and safely ignored."',
            ],
            [
                "Given a valid signed webhook is received, when schema validation passes, then the system stores and processes it once.",
                "Given an unsigned webhook is received, when validation runs, then the payload is rejected and logged without domain side effects.",
                "Given a downstream system is unavailable, when delivery fails, then the system retries and surfaces the failure on the integration dashboard.",
            ],
        )

        add_requirement(
            doc,
            f"SRS-SFL-{sid}-05",
            f"Expose {system['short']} Dashboards and Reports",
            "Facilities, Security, HSE or Logistics Manager",
            f"monitor {system['short']} indicators and drill into exceptions",
            "management can act quickly on readiness, safety, compliance and service risks",
            [
                f"The system shall expose dashboard indicators for {system['dashboard']}.",
                "Dashboard data shall be filterable by site, date range, status, priority, owner and operating mode where applicable.",
                "Dashboard records shall link back to source workflows and evidence where the user has permission.",
                "Dashboard snapshots shall be suitable for operational review and go-live readiness reporting.",
            ],
            [
                "Dashboard snapshot ID, metric code, period, site scope, generated timestamp and source record references.",
            ],
            [
                "Users may only view dashboard data for authorised sites and roles.",
                "Dashboard counts must reconcile to source workflow/read-model records.",
                "Critical safety and examination-readiness indicators must display stale-data warnings where freshness thresholds are breached.",
            ],
            [
                "Manager opens the dashboard.",
                "System applies role and site filters.",
                "User reviews indicators and drills into exceptions.",
                "System records dashboard access where required for sensitive views.",
            ],
            [
                'Data Stale - "Dashboard data is older than the configured freshness threshold."',
                'No Scope - "No site scope is assigned to your user profile."',
                'Restricted Drilldown - "You do not have permission to view the underlying record."',
            ],
            [
                "Given a manager opens the dashboard, when they have site scope, then the system displays only authorised indicators.",
                "Given a dashboard metric is stale, when it renders, then the system shows a stale-data warning.",
                "Given a user clicks an exception, when they have permission, then the system opens the source workflow or evidence record.",
            ],
        )


def add_data_interfaces_nfr(doc: Document) -> None:
    doc.add_heading("21. Data Requirements", level=1)
    doc.add_heading("21.1 Key Entities and Relationships", level=2)
    add_table(doc, ["Entity", "Related Entities", "Relationship"], [
        ["Site / Building / Room / Zone", "Facility Asset, Booking, Work Order, Access Zone, Device", "One site contains many buildings, rooms, zones, devices and operational records."],
        ["Facility Asset", "Work Order, Preventive Schedule, Evidence", "One asset may have many faults, work orders, schedules and closure evidence records."],
        ["Room Booking", "Room, Resource, Host, Setup Task, Security Task", "One booking reserves one or more resources and may create setup/security/logistics tasks."],
        ["Visitor Visit", "Visitor, Host, Badge, Access Zone, Roll-call", "One visit belongs to one visitor and host; visit status feeds evacuation roll-call."],
        ["Access Event", "Credential, Reader, Door, Person/Visitor, Incident", "Many access events occur at one reader and may create incidents or exceptions."],
        ["Camera / CCTV Event", "Camera, Incident, Evidence Request, Export Approval", "Camera health and evidence requests link to incident and export governance records."],
        ["Alarm / Life-Safety Event", "Zone, SOC Queue, Emergency Notification, CMMS Work Order", "Events may trigger response workflow, notifications and maintenance tickets."],
        ["HSE Incident", "Reporter, Location, Evidence, RCA, CAPA", "One incident may have many evidence items, witness notes and corrective actions."],
        ["Vehicle", "Driver, Assignment, Fuel Transaction, Service Record", "One vehicle has many assignments, fuel entries, service records and compliance documents."],
        ["Fuel Transaction", "Vehicle, Driver, Receipt, Anomaly Case", "Fuel transactions reconcile against limits, odometer readings and driver logbooks."],
        ["Courier Item / Dispatch Manifest", "Custody Handoff, Receipt, Exception, Evidence", "One manifest contains many items and custody events."],
        ["Emergency Notification", "Template, Zone, Channel, Recipient, Acknowledgement", "One notification may be sent through many channels to many recipients."],
        ["Audit Event / Evidence", "All operational entities", "Every state-changing action creates audit; evidence records carry retention and hash metadata."],
    ], [1.65, 2.2, 2.65])

    doc.add_heading("21.2 Data Governance Rules", level=2)
    add_bullets(doc, [
        "All operational records must carry site scope, owning platform, creator, timestamps and audit correlation ID.",
        "Biometric, CCTV, visitor, access, HSE and dispatch evidence must have data classification and retention class.",
        "Events must publish references and metadata, not sensitive content, unless a formally approved secure integration requires it.",
        "Data imported from external systems must retain source system, source identifier, ingestion timestamp and validation status.",
        "Records used for examination continuity or dispatch chain-of-custody must be protected from deletion and governed by retention/legal-hold rules.",
    ])

    doc.add_heading("22. External Interface Requirements", level=1)
    add_table(doc, ["External System", "Interface Type", "Direction", "Purpose"], [
        ["Enterprise Identity Provider", "OIDC/OAuth2/JWT/JWKS; admin adapter", "Inbound/Outbound", "Authentication, role/site-scope claims, service accounts and provisioning where approved."],
        ["HRMS", "REST API / events", "Inbound", "Staff, host, driver, joiner-mover-leaver and organisational reference data."],
        ["NBES / Examination Operations", "REST API / events", "Inbound/Outbound", "Examination centres, schedules, examination mode, hall readiness and dispatch context."],
        ["CCTV / VMS / NVR", "Vendor API / webhook / controlled export", "Inbound/Outbound", "Camera health, recording status, incident linkage and evidence export references."],
        ["Access Control / Biometric Readers", "Vendor API / event stream", "Inbound/Outbound", "Access granted/denied events, credential rules, reader health and overrides."],
        ["Intrusion Alarm Panels", "Webhook / gateway / monitoring feed", "Inbound", "Zone alarm, tamper, acknowledgement and restoration events."],
        ["Fire / Life-Safety Panels", "Panel gateway / BMS bridge / monitoring feed", "Inbound", "Fire, smoke, panic, fault and inspection events; SFL observes and supplements only."],
        ["Notification Providers", "SMS/email/push/voice/siren/signage APIs", "Outbound/Inbound receipts", "Emergency and operational notifications, delivery receipts and acknowledgements."],
        ["Finance / Procurement / Vendor Master", "REST API / file import", "Inbound/Outbound", "Vendor, fuel, payment, procurement and contract references."],
        ["Fleet/Fuel Provider Systems", "API / file import / mobile capture", "Inbound", "Fuel issues, receipts, odometer, limits and anomalies."],
        ["Document Management / Evidence Store", "REST API / object store adapter", "Outbound", "Evidence files, export packages, document references and retention holds."],
        ["SIEM / SOC / NOC Monitoring", "Event stream / syslog / API", "Outbound", "Security, integration health, device health and audit alerts."],
        ["Analytics / Reporting", "Read model / API / event stream", "Outbound", "Dashboards, KPIs, readiness scoring and management reporting."],
    ], [1.75, 1.55, 1.1, 2.1])

    doc.add_heading("23. Non-Functional Requirements", level=1)
    doc.add_heading("23.1 Security", level=2)
    add_bullets(doc, [
        "Authentication, MFA, password and session lifecycle are delegated to the enterprise identity provider.",
        "All APIs require role and site-scope authorization. Privileged actions require explicit permissions and audit logging.",
        "Inbound integrations require HMAC or mTLS, source allowlisting, schema validation and idempotency before domain processing.",
        "Sensitive data is encrypted at rest and in transit. Secrets are stored in approved secret-management facilities, not source code.",
        "Break-glass emergency actions are permitted only for approved roles and templates, and are audited after the fact.",
    ])
    doc.add_heading("23.2 Performance and Scalability", level=2)
    add_bullets(doc, [
        "Routine workflow pages should respond within operationally acceptable browser response times under normal load.",
        "Device and webhook ingestion must process bursts without losing events; queued processing is acceptable where domain action is not life-safety critical.",
        "Emergency fast-lane notifications should meet the approved latency target once final provider and infrastructure numbers are confirmed.",
        "Dashboards must show stale-data warnings when freshness thresholds are breached.",
    ])
    doc.add_heading("23.3 Availability and Reliability", level=2)
    add_bullets(doc, [
        "Core platform functions must support backup, restore, health checks and monitored deployment across development, staging and production.",
        "Examination-centre edge runtime must support local capture for approved edge-authoritative workflows during WAN loss and reconcile through outbox on restore.",
        "Failed outbound integrations must retry with backoff and escalate to operations after configured attempts.",
        "Platform mode changes, such as Routine to Examination Mode, must be explicit, audited and reversible only by authorised roles.",
    ])
    doc.add_heading("23.4 Usability and Accessibility", level=2)
    add_bullets(doc, [
        "The SFL portal shall use a consistent design system across facilities, safety, fleet and logistics modules.",
        "Field workflows shall be usable on mobile browsers for evidence capture, inspections, dispatch receipts and incident reporting.",
        "Forms shall preserve partially entered data where possible to avoid loss during field connectivity interruptions.",
        "Interfaces should meet recognised accessibility practices for keyboard navigation, readable contrast and clear error messaging.",
    ])
    doc.add_heading("23.5 Audit, Logging and Traceability", level=2)
    add_bullets(doc, [
        "Every state-changing action must capture actor, timestamp, action, source channel, correlation ID and before/after values where applicable.",
        "Audit records must be append-only and tamper-evident through hash-chain or equivalent replay-verifiable design.",
        "Evidence access and export events must be audited with reason, approver, recipient and retention/legal-hold status.",
        "OpenTelemetry trace context must propagate through HTTP, outbox and event processing paths.",
    ])
    doc.add_heading("23.6 Data Retention and Privacy", level=2)
    add_bullets(doc, [
        "Visitor, biometric, CCTV, incident, access and dispatch records must be classified and assigned approved retention schedules.",
        "Data minimisation applies to published events and reports; sensitive payloads are not exposed where metadata references are sufficient.",
        "Data-subject request support must allow authorised retrieval, correction workflow and deletion/retention decision where legally permitted.",
        "CCTV and biometric processing requires documented lawful basis, DPIA where required and controlled evidence-release workflow.",
    ])
    doc.add_heading("23.7 Localisation", level=2)
    add_bullets(doc, [
        "Dates, times, phone numbers, addresses and official labels must support Ghana operational context.",
        "Emergency templates must support approved message text per site, channel and scenario.",
        "System configuration should allow terminology and labels to be adjusted without code changes where operationally necessary.",
    ])
    doc.add_heading("23.8 Configurability and Maintainability", level=2)
    add_bullets(doc, [
        "Operational thresholds, SLAs, escalations, roles, zones, readiness checklists, fuel limits and retention schedules must be runtime-configurable and versioned.",
        "Spring Boot modules must isolate infrastructure adapters from domain/application code using ports and adapters.",
        "Architecture tests must enforce module boundaries, contracts-only references, no cross-schema foreign keys and no provider names outside adapters.",
    ])


def add_appendix(doc: Document) -> None:
    doc.add_heading("Appendix A. Phase and Source Traceability", level=1)
    doc.add_paragraph("The following table preserves the authoritative Cluster 9 system list from the parent mapping. Phase 1 implementation is limited to the 13 Fast-Track systems; Phase 2 and Phase 3 systems are reserved for future scope unless specifically called out as supporting references.")
    add_table(doc, ["System ID", "System Name", "Mapping Phase"], [[a, b, c] for a, b, c in FULL_CLUSTER], [0.85, 4.55, 1.1])

    doc.add_heading("Appendix B. Phase 1 Release Alignment", level=1)
    add_table(doc, ["Release", "Platform / Focus", "Systems", "Exit Criteria"], [
        ["Release 0", "Foundation", "Identity ports, audit/evidence, workflow, integration hub, cache, messaging, edge skeleton and architecture tests", "Infrastructure ports and adapters are proven before operational modules depend on them."],
        ["Release 1", "Facilities Core - SFL.IFIMP", "S152, S153, S159", "Facilities master data, work orders and booking create auditable workflows and readiness indicators."],
        ["Release 2", "Safety and Security Core - SFL.SSEMP", "S160, S160a, S161, S162, S162a, S163, S174", "Visitor, access, CCTV, alarm, life-safety, HSE and emergency notification workflows are operational and audited."],
        ["Release 3", "Fleet and Logistics Core - SFL.FTLMP", "S166, S168_fuel, S171", "Fleet, fuel/logbook and courier/dispatch workflows preserve custody and exception handling."],
        ["Release 4", "Examination and Operational Readiness", "Cross-platform readiness, dashboards and edge operation", "Hall readiness, emergency response, dispatch traceability and edge failover are commissioned."],
        ["Release 5", "Commissioning and Go-Live", "All Phase 1 systems", "Acceptance tests, training, SOPs, data load, integrations and management sign-off are complete."],
    ], [0.8, 1.55, 2.0, 2.15])


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = setup_document()
    add_cover(doc)
    add_front_matter(doc)
    add_intro(doc)
    add_overall(doc)
    add_system_requirements(doc)
    add_data_interfaces_nfr(doc)
    add_appendix(doc)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
