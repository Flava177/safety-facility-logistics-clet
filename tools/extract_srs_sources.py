from __future__ import annotations

import argparse
import re
from pathlib import Path

import pdfplumber
from docx import Document


def extract_docx(path: Path) -> tuple[str, list[str]]:
    doc = Document(path)
    lines: list[str] = []
    headings: list[str] = []
    for paragraph in doc.paragraphs:
        text = " ".join(paragraph.text.split())
        if not text:
            continue
        style = paragraph.style.name if paragraph.style is not None else ""
        if style.startswith("Heading") or re.match(r"^\d+(\.\d+)*\s+", text):
            headings.append(text)
        lines.append(text)

    for table_number, table in enumerate(doc.tables, 1):
        lines.append(f"[TABLE {table_number}]")
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    return "\n".join(lines), headings


def extract_pdf(path: Path) -> tuple[str, list[str]]:
    pages: list[str] = []
    headings: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, 1):
            text = page.extract_text(x_tolerance=1, y_tolerance=3) or ""
            pages.append(f"\n\n=== PAGE {page_number} ===\n{text}")
            for line in text.splitlines():
                clean = " ".join(line.split())
                if re.match(r"^(\d+(\.\d+)*\.?\s+|[A-Z][A-Z0-9 /&(),:-]{8,}$)", clean):
                    headings.append(clean)
    return "\n".join(pages), headings


def extract_md(path: Path) -> tuple[str, list[str]]:
    text = path.read_text(encoding="utf-8", errors="replace")
    headings = [line.strip("# ").strip() for line in text.splitlines() if line.startswith("#")]
    return text, headings


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--source-dir", required=True)
    parser.add_argument("--out-dir", required=True)
    args = parser.parse_args()

    source_dir = Path(args.source_dir)
    out_dir = Path(args.out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    manifest: list[str] = []
    for path in sorted(source_dir.iterdir()):
        if path.suffix.lower() not in {".docx", ".pdf", ".md"}:
            continue
        if path.suffix.lower() == ".docx":
            text, headings = extract_docx(path)
        elif path.suffix.lower() == ".pdf":
            text, headings = extract_pdf(path)
        else:
            text, headings = extract_md(path)

        stem = path.stem
        text_path = out_dir / f"{stem}.txt"
        headings_path = out_dir / f"{stem}.headings.txt"
        text_path.write_text(text, encoding="utf-8")
        headings_path.write_text("\n".join(headings), encoding="utf-8")
        manifest.append(f"{path.name}\t{text_path.name}\t{headings_path.name}\t{len(text):,} chars")

    (out_dir / "manifest.tsv").write_text("\n".join(manifest), encoding="utf-8")


if __name__ == "__main__":
    main()
